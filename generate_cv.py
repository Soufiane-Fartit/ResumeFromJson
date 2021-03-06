from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from data import *

document = Document()

def add_header(doc, header):
    for i, elem in enumerate(header) :
        if i == 0 :
            p = doc.add_paragraph(elem)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 0
            p.paragraph_format.space_before = Cm(0)
            p.paragraph_format.space_after = Cm(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            for run in p.runs :
                run.font.size = Pt(14)
        else :
            p = doc.add_paragraph(elem)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 0
            p.paragraph_format.space_before = Cm(0)
            p.paragraph_format.space_after = Cm(0.7)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    return doc


def add_profil(doc, profil):
    doc = add_title(doc, "PROFIL")
    for elem in profil :
        p = doc.add_paragraph(elem)
        p.paragraph_format.line_spacing = 1

    return doc

def add_formation(doc, formation):
    doc = add_title(doc, "FORMATION ACADEMIQUE")
    for form in formation :
        p = doc.add_paragraph(form['date']+' : '+form['diplome']+' - '+form['ecole'])
        p.paragraph_format.line_spacing = 0
        p.paragraph_format.space_before = Cm(0)
        p.paragraph_format.space_after = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    return doc

def add_experience(doc, experience):
    doc = add_title(doc, "EXPERIENCE PROFESSIONNELLE")
    for exp in experience :
        # ADD JOB TITLE LINE
        p = doc.add_paragraph(exp['title'])
        p.paragraph_format.space_before = Cm(0.2)
        p.paragraph_format.space_after = Cm(0)
        p.paragraph_format.line_spacing = 1
        for run in p.runs:
            run.bold = True
        
        # ADD COMPANY NAME, ADRESS, DATES
        p = doc.add_paragraph(exp['company'] + ', '+ (9*"\t").join([exp['site'], exp['date']]))
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_before = Cm(0) 
        p.paragraph_format.space_after = Cm(0) 
        for run in p.runs:
            run.italic = True
        
        # ADD CONTENT
        if exp['text'] != "" :
            p = doc.add_paragraph(exp['text'])
            p.paragraph_format.line_spacing = 1
        
        # ADD BULLETPOINTS
        for task in exp['tasks'] :
            p = doc.add_paragraph(task, style='List Bullet')
            p.paragraph_format.line_spacing = 1
            p.paragraph_format.space_before = Cm(0)
            p.paragraph_format.space_after = Cm(0)
        
        doc.add_paragraph('Environnement : ' + ', '.join(exp['Environnement']))

    return doc

def add_skills(doc, skills):
    doc = add_title(doc, "COMPETENCES TECHNIQUES")
    for k, v in skills.items():
        if len(v) !=0 :
            p = doc.add_paragraph(k + ' : ' + ', '.join(v))
            p.paragraph_format.line_spacing = 1
            p.paragraph_format.space_before = Cm(0)
            p.paragraph_format.space_after = Cm(0)

    return doc

def add_certif_projects(doc, certifs, projects):
    doc = add_title(doc, "CERTIFICATIONS ET PROJETS PERSONNELS")

    for certif in certifs:
        p = doc.add_paragraph(certif, style='List Bullet')
        p.paragraph_format.line_spacing = 1
    for project in projects:
        p = doc.add_paragraph(project, style='List Bullet')
        p.paragraph_format.line_spacing = 1

    return doc

def add_langues(doc, langues):
    doc = add_title(doc, "LANGUES")
    
    p = doc.add_paragraph((5*' '+3*'\t').join([k+' : '+v for k,v in langues.items()]))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_before = Cm(0) 
    p.paragraph_format.space_after = Cm(0)

    return doc

def set_margin(doc, margin = 1.0):
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin) 
    
    return doc

def first_child_found_in(parent, tagnames):
    """
    Return the first child of parent with tag in *tagnames*, or None if
    not found.
    """
    for tagname in tagnames:
        child = parent.find(qn(tagname))
        if child is not None:
            return child
    return None

def insert_element_before(parent, elm, successors):
    """
    Insert *elm* as child of *parent* before any existing child having
    tag name found in *successors*.
    """
    successor = first_child_found_in(parent, successors)
    if successor is not None:
        successor.addprevious(elm)
    else:
        parent.append(elm)
    return elm


def add_title(doc, title):
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Cm(0.3)
    pp.paragraph_format.space_after = Cm(0.1)
    run = pp.add_run(title)
    #run.bold = True
    """

    # ADD COLOR SHADING

    # Get the XML tag
    tag = run._r

    # Create XML element
    shd = OxmlElement('w:shd')

    # Add attributes to the element
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '90a4ae')

    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)
    tag.rPr.append(shd)
    """
    p = pp._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    insert_element_before(pPr, pBdr, successors=(
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    ))
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

    return doc

def redact_cv(doc, header, profil, formation, experience, skills, certifs, projects, langues):
    # ADD NAME, PHONE, ADRESS, EMAIL ...
    doc = add_header(doc, header)

    # ADD A FEW LINES ABOUT ME
    doc = add_profil(doc, profil)

    # ADD EDUCATION SECTION
    doc = add_formation(doc, formation)

    # ADD PRO EXP SECTION
    doc = add_experience(doc, experience)

    # ADD SKILLS SECTION
    doc = add_skills(doc, skills)

    # ADD CERTIFS PERS PROJECTS
    doc = add_certif_projects(doc, certifs, projects)

    # ADD SPOKE LANGUAGES
    doc = add_langues(doc, langues)

    # SET MARGINS
    doc = set_margin(doc)

    return doc

document = redact_cv(document, header, profil, formation, experience, skills, certifs, projects, langues)

document.save('output.docx')
